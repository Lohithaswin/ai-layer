from __future__ import annotations

import docx
from pathlib import Path

from src.config import DOCS_DIR, CHILD_CHUNK_SIZE, CHILD_CHUNK_OVERLAP, PARENT_MAX_CHARS
from src.pdf_loader import _split_child_chunks, _build_child_parent_text

def load_word_chunks(file_path: Path, metadata: dict | None = None) -> list[dict]:
    """
    Load Word documents and convert into chunks for RAG.
    """
    try:
        rel_path = str(file_path.relative_to(DOCS_DIR.parent))
    except Exception:
        try:
            rel_path = str(file_path.relative_to(file_path.parents[1]))
        except Exception:
            rel_path = file_path.name
    rel_path = rel_path.replace("\\", "/")

    chunks: list[dict] = []
    chunk_index = 0

    try:
        doc = docx.Document(file_path)
        
        # Simple extraction: all paragraphs and tables sequentially
        # For a more robust extraction, we could parse XML to keep them in exact order.
        # But for many documents, just getting paragraphs then tables is acceptable,
        # or we iterate through document.element.body to preserve order.
        
        parts = []
        for element in doc.element.body:
            if element.tag.endswith('p'):
                # It's a paragraph
                for p in doc.paragraphs:
                    if p._element == element:
                        text = p.text.strip()
                        if text:
                            parts.append(text)
                        break
            elif element.tag.endswith('tbl'):
                # It's a table
                for t in doc.tables:
                    if t._element == element:
                        for row in t.rows:
                            row_data = []
                            for cell in row.cells:
                                c_text = cell.text.strip().replace("\n", " ")
                                if c_text and c_text not in row_data:  # naive merge cell handling
                                    row_data.append(c_text)
                            if row_data:
                                parts.append(" | ".join(row_data))
                        break
                        
        page_text = "\n\n".join(parts)
        
        if not page_text.strip():
            return []

        # Word docs don't have clear page boundaries with python-docx easily,
        # so we treat the whole document as "Page 1".
        parent_id = f"{rel_path}|1"
        
        for piece, start, end in _split_child_chunks(
            page_text, CHILD_CHUNK_SIZE, CHILD_CHUNK_OVERLAP
        ):
            row_dict = {
                "text": piece,
                "parent_text": _build_child_parent_text(
                    page_text, start, end, PARENT_MAX_CHARS
                ),
                "parent_id": parent_id,
                "source_file": rel_path,
                "page": 1,
                "chunk_index": chunk_index,
                "section_title": "Document",
            }
            if metadata:
                row_dict.update(metadata)
            chunks.append(row_dict)
            chunk_index += 1
            
    except Exception as e:
        print(f"Error reading Word document {file_path}: {e}")

    return chunks
